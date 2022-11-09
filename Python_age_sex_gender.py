# Imports
import urllib.request, json
from docx import Document

# List of FED IDs
fed_id_list =['2013A000435001']

# Iterate over FED IDs
for fed_id in fed_id_list:

        # Get FED 2016 Census Data
        with urllib.request.urlopen('https://www12.statcan.gc.ca/rest/census-recensement/CPR2016.json?lang=E&dguid='+fed_id) as url:
                fed_data = json.loads(url.read().decode())
        fed_name = fed_data['DATA'][0][4]
        
        # Get FED 2021 Census Data (including sex, gender, type of dwelling)
        with urllib.request.urlopen('https://api.statcan.gc.ca/census-recensement/profile/sdmx/rest/data/STC_CP,DF_FED/A5.'+fed_id+
                                '.1.1.1?format=jsondata&detail=dataonly') as url:
                fed_data_2021 = json.loads(url.read().decode())
        total_pop_count_2021 = float(fed_data_2021['data']['dataSets'][0]['series']['0:0:0:0:0']['observations']['0'][0])
        
        with urllib.request.urlopen('https://api.statcan.gc.ca/census-recensement/profile/sdmx/rest/data/STC_CP,DF_FED/A5.'+fed_id+
                                '.1.3.1?format=jsondata&detail=dataonly') as url:
                fed_data_2021 = json.loads(url.read().decode())
        total_pop_growth_rate_2016to2021 = float(fed_data_2021['data']['dataSets'][0]['series']['0:0:0:0:0']['observations']['0'][0])

        with urllib.request.urlopen('https://api.statcan.gc.ca/census-recensement/profile/sdmx/rest/data/STC_CP,DF_FED/A5.'+fed_id+
                                '.1.8.1?format=jsondata&detail=dataonly') as url:
                fed_data_2021 = json.loads(url.read().decode())
        total_age_population = float(fed_data_2021['data']['dataSets'][0]['series']['0:0:0:0:0']['observations']['0'][0])

        with urllib.request.urlopen('https://api.statcan.gc.ca/census-recensement/profile/sdmx/rest/data/STC_CP,DF_FED/A5.'+fed_id+
                                '.2.8.1?format=jsondata&detail=dataonly') as url:
                fed_data_2021 = json.loads(url.read().decode())
        total_age_population_male = float(fed_data_2021['data']['dataSets'][0]['series']['0:0:0:0:0']['observations']['0'][0])

        with urllib.request.urlopen('https://api.statcan.gc.ca/census-recensement/profile/sdmx/rest/data/STC_CP,DF_FED/A5.'+fed_id+
                                '.3.8.1?format=jsondata&detail=dataonly') as url:
                fed_data_2021 = json.loads(url.read().decode())
        total_age_population_female = float(fed_data_2021['data']['dataSets'][0]['series']['0:0:0:0:0']['observations']['0'][0])

        with urllib.request.urlopen('https://api.statcan.gc.ca/census-recensement/profile/sdmx/rest/data/STC_CP,DF_FED/A5.'+fed_id+
                                '.1.4.1?format=jsondata&detail=dataonly') as url:
                fed_data_2021 = json.loads(url.read().decode())
        total_private_dwellings = float(fed_data_2021['data']['dataSets'][0]['series']['0:0:0:0:0']['observations']['0'][0])

 # Generate DOCX Document
        document = Document()

        document.add_heading(f'Profile:')

        document.add_heading(f'Constituency: {fed_name}, Ontario')

        document.add_heading('Gender:')
        document.add_paragraph(
                        f'According to the 2021 Census, the population of the {fed_name} FED was {total_pop_count_2021:,.0f}, '+
                        ('an increase ' if total_pop_growth_rate_2016to2021>0 else 'a decrease ')+
                        f'of {total_pop_growth_rate_2016to2021/100:.1%} from 2016.  '+
                        f'This is in contrast to the province where the population grew by 775,448 (5.8%). '+
                        f'According to the 2021 Census, {total_age_population_male:,.0f} ({total_age_population_male/total_age_population:.1%}) of '+
                        f'the population were male, while {total_age_population_female:,.0f} ({total_age_population_female/total_age_population:.1%}) were female.')

        document.add_heading('Household and dwelling characteristics:')
        document.add_paragraph(
                        f'According to the 2021 Census, the number of private dwellings in the {fed_name} FED was {total_private_dwellings:,.0f}.')

        document.save(fed_name +'.docx')
