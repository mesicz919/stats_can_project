#!/usr/bin/env python
# coding: utf-8

# In[1]:


import urllib.request, json

import requests
import pandas as pd
import io
import json
import os
import zipfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta


# In[2]:



#define function for API to download statcan tables
def StatcanTable(TableNumber):
    TableNumber = str(TableNumber)
    r = requests.get("https://www150.statcan.gc.ca/n1/en/tbl/csv/" + TableNumber + "-eng.zip", stream=True)
    print(r)
    z = zipfile.ZipFile(io.BytesIO(r.content))
    z.extractall(os.getcwd())
    RawData = pd.read_csv(TableNumber + ".csv",low_memory=False)
    return RawData


# In[3]:



# List of FED IDs
fed_id_list =['2013A000435001']#,'2013A000435002','2013A000435003','2013A000435004','2013A000435005','2013A000435006','2013A000435007','2013A000435008','2013A000435009','2013A000435010','2013A000435011','2013A000435012','2013A000435013','2013A000435014','2013A000435015','2013A000435016','2013A000435017','2013A000435018','2013A000435019','2013A000435020','2013A000435021','2013A000435022','2013A000435023','2013A000435024','2013A000435025','2013A000435026','2013A000435027','2013A000435028','2013A000435029','2013A000435030','2013A000435031','2013A000435032','2013A000435033','2013A000435034','2013A000435035','2013A000435036','2013A000435037','2013A000435038','2013A000435039','2013A000435040','2013A000435041','2013A000435042','2013A000435043','2013A000435044','2013A000435045','2013A000435046','2013A000435047','2013A000435048','2013A000435049','2013A000435050','2013A000435051','2013A000435052','2013A000435053','2013A000435054','2013A000435055','2013A000435056','2013A000435057','2013A000435058','2013A000435059','2013A000435060','2013A000435061','2013A000435062','2013A000435063','2013A000435064','2013A000435065','2013A000435066','2013A000435067','2013A000435068','2013A000435069','2013A000435070','2013A000435071','2013A000435072','2013A000435073','2013A000435074','2013A000435075','2013A000435076','2013A000435077','2013A000435078','2013A000435079','2013A000435080','2013A000435081','2013A000435082','2013A000435083','2013A000435084','2013A000435085','2013A000435086','2013A000435087','2013A000435088','2013A000435089','2013A000435090','2013A000435091','2013A000435092','2013A000435093','2013A000435094','2013A000435095','2013A000435096','2013A000435097','2013A000435098','2013A000435099','2013A000435100','2013A000435101','2013A000435102','2013A000435103','2013A000435104','2013A000435105','2013A000435106','2013A000435107','2013A000435108','2013A000435109','2013A000435110','2013A000435111','2013A000435112','2013A000435113','2013A000435114','2013A000435115','2013A000435116','2013A000435117','2013A000435118','2013A000435119','2013A000435120','2013A000435121']


# In[4]:




#### Add Census 2021 Income Data ############


# Iterate over FED IDs
for fed_id in fed_id_list:

        # Get FED 2016 Census Data
        with urllib.request.urlopen('https://www12.statcan.gc.ca/rest/census-recensement/CPR2016.json?lang=E&dguid='+fed_id) as url:
                fed_data = json.loads(url.read().decode())
        fed_name = fed_data['DATA'][0][4]
        total_pop_count = fed_data['DATA'][0][13]
        immigrants_total = fed_data['DATA'][1139][13]
        immigrants_count = fed_data['DATA'][1141][13]
        immigrants_2011to2016_count = fed_data['DATA'][1148][13]
        non_citizen_count = fed_data['DATA'][1138][13]
        indigenous_count = fed_data['DATA'][1289][13]
        vis_minority_count = fed_data['DATA'][1323][13]
        vis_minority_total =  fed_data['DATA'][1322][13]
        vis_minority_groups = [
                        ['South Asian',fed_data['DATA'][1324][13]],
                        ['Chinese',fed_data['DATA'][1325][13]],
                        ['Black',fed_data['DATA'][1326][13]],
                        ['Filipino',fed_data['DATA'][1327][13]],
                        ['Latin American',fed_data['DATA'][1328][13]],
                        ['Arab',fed_data['DATA'][1329][13]],
                        ['Southeast Asian',fed_data['DATA'][1330][13]],
                        ['West Asian',fed_data['DATA'][1331][13]],
                        ['Korean',fed_data['DATA'][1332][13]],
                        ['Japanese',fed_data['DATA'][1333][13]],
                        ['Multiracial',fed_data['DATA'][1335][13]]
                        ]
        vis_minority_groups.sort(key= lambda x:x[1], reverse=True)
        industry_total =  fed_data['DATA'][1898][13]
        industry_groups = [
                        ['agriculture, forestry, fishing and hunting',fed_data['DATA'][1899][13]],
                        ['mining, quarrying, and oil and gas extraction',fed_data['DATA'][1900][13]],
                        ['utilities',fed_data['DATA'][1901][13]],
                        ['construction',fed_data['DATA'][1902][13]],
                        ['manufacturing',fed_data['DATA'][1903][13]],
                        ['wholesale trade',fed_data['DATA'][1904][13]],
                        ['retail trade',fed_data['DATA'][1905][13]],
                        ['transportation and warehousing',fed_data['DATA'][1906][13]],
                        ['information and cultural industries',fed_data['DATA'][1907][13]],
                        ['finance and insurance',fed_data['DATA'][1908][13]],
                        ['real estate and rental and leasing',fed_data['DATA'][1909][13]],
                        ['professional, scientific and technical services',fed_data['DATA'][1910][13]],
                        ['management of companies and enterprises',fed_data['DATA'][1911][13]],
                        ['administrative and support, waste management and remediation services',fed_data['DATA'][1912][13]],
                        ['educational services',fed_data['DATA'][1913][13]],
                        ['health care and social assistance',fed_data['DATA'][1914][13]],
                        ['arts, entertainment and recreation',fed_data['DATA'][1915][13]],
                        ['accommodation and food services',fed_data['DATA'][1916][13]],
                        ['other services (except public administration)',fed_data['DATA'][1917][13]],
                        ['public administration',fed_data['DATA'][1918][13]]
                        ]
        industry_groups.sort(key= lambda x:x[1], reverse=True)
        englishandfrench = fed_data['DATA'][102][13]        
        language_total = fed_data['DATA'][99][13]    
        language_groups = [
                        ['Mandarin',fed_data['DATA'][617][13]],
                        ['Cantonese',fed_data['DATA'][615][13]],
                        ['Punjabi(Panjabi)',fed_data['DATA'][567][13]],
                        ['Spanish',fed_data['DATA'][581][13]],
                        ['Arabic',fed_data['DATA'][481][13]],
                        ['Urdu',fed_data['DATA'][570][13]],
                        ['Tamil',fed_data['DATA'][514][13]],
                        ['Persian(Farsi)',fed_data['DATA'][574][13]],
                        ['Tagalog(PilipinoFilipino)',fed_data['DATA'][504][13]],
                        ['Italian',fed_data['DATA'][578][13]],
                        ['Portuguese',fed_data['DATA'][579][13]],
                        ['Russian',fed_data['DATA'][533][13]],
                        ['Polish',fed_data['DATA'][532][13]],
                        ['Korean',fed_data['DATA'][586][13]],
                        ['Vietnamese',fed_data['DATA'][492][13]],
                        ['EnglishandFrench',fed_data['DATA'][645][13]],
                        ['Gujarati',fed_data['DATA'][560][13]],
                        ['German',fed_data['DATA'][549][13]],
                        ['Bengali',fed_data['DATA'][559][13]],
                        ['Hindi',fed_data['DATA'][561][13]],
                        ['Serbian',fed_data['DATA'][534][13]],
                        ['Greek',fed_data['DATA'][556][13]],
                        ['Romanian',fed_data['DATA'][580][13]],
                        ['Ukrainian',fed_data['DATA'][538][13]],
                        ['Somali',fed_data['DATA'][477][13]],
                        ['Hungarian',fed_data['DATA'][641][13]],
                        ['Croatian',fed_data['DATA'][529][13]],
                        ['Turkish',fed_data['DATA'][634][13]],
                        ['Chinesen.o.s.',fed_data['DATA'][621][13]],
                        ['AssyrianNeo-Aramaic',fed_data['DATA'][482][13]],
                        ['Albanian',fed_data['DATA'][519][13]],
                        ['Malayalam',fed_data['DATA'][513][13]],
                        ['Armenian',fed_data['DATA'][520][13]],
                        ['Pashto',fed_data['DATA'][573][13]],
                        ['Frenchandnon-officiallanguage',fed_data['DATA'][647][13]],
                        ['Macedonian',fed_data['DATA'][531][13]],
                        ['Japanese',fed_data['DATA'][583][13]],
                        ['MinNan(ChaochowTeochowFukienTaiwanese)',fed_data['DATA'][619][13]],
                        ['Telugu',fed_data['DATA'][515][13]],
                        ['Nepali',fed_data['DATA'][565][13]],
                        ['Amharic',fed_data['DATA'][480][13]],
                        ['Hebrew',fed_data['DATA'][485][13]],
                        ['Bulgarian',fed_data['DATA'][528][13]],
                        ['Sinhala(Sinhalese)',fed_data['DATA'][569][13]],
                        ['Dutch',fed_data['DATA'][547][13]],
                        ['Akan(Twi)',fed_data['DATA'][590][13]],
                        ['Ilocano',fed_data['DATA'][499][13]],
                        ['Creolen.o.s.',fed_data['DATA'][509][13]],
                        ['Tigrigna',fed_data['DATA'][487][13]],
                        ['Tibetan',fed_data['DATA'][626][13]],
                        ['Khmer(Cambodian)',fed_data['DATA'][491][13]],
                        ['Wu(Shanghainese)',fed_data['DATA'][620][13]],
                        ['Kurdish',fed_data['DATA'][572][13]],
                        ['Slovak',fed_data['DATA'][536][13]],
                        ['Ojibway',fed_data['DATA'][405][13]],
                        ['Bosnian',fed_data['DATA'][527][13]],
                        ['Czech',fed_data['DATA'][530][13]],
                        ['ChaldeanNeo-Aramaic',fed_data['DATA'][483][13]],
                        ['Oji-Cree',fed_data['DATA'][406][13]],
                        ['Sindhi',fed_data['DATA'][568][13]],
                        ['Hakka',fed_data['DATA'][616][13]],
                        ['Lao',fed_data['DATA'][629][13]],
                        ['Marathi',fed_data['DATA'][564][13]],
                        ['Cebuano',fed_data['DATA'][496][13]],
                        ['Creen.o.s.',fed_data['DATA'][399][13]],
                        ['Swahili',fed_data['DATA'][602][13]],
                        ['Serbo-Croatian',fed_data['DATA'][535][13]],
                        ['Finnish',fed_data['DATA'][640][13]],
                        ['Malay',fed_data['DATA'][501][13]],
                        ['AmericanSignLanguage',fed_data['DATA'][610][13]],
                        ['Slovene(Slovenian)',fed_data['DATA'][537][13]],
                        ['Lithuanian',fed_data['DATA'][524][13]],
                        ['Thai',fed_data['DATA'][630][13]],
                        ['Yoruba',fed_data['DATA'][604][13]],
                        ['Maltese',fed_data['DATA'][486][13]],
                        ['Kannada',fed_data['DATA'][512][13]],
                        ['Burmese',fed_data['DATA'][624][13]],
                        ['Hiligaynon',fed_data['DATA'][498][13]],
                        ['Latvian',fed_data['DATA'][523][13]],
                        ['Oromo',fed_data['DATA'][476][13]],
                        ['Estonian',fed_data['DATA'][639][13]],
                        ['Georgian',fed_data['DATA'][585][13]],
                        ['Azerbaijani',fed_data['DATA'][633][13]],
                        ['Afrikaans',fed_data['DATA'][545][13]],
                        ['Uzbek',fed_data['DATA'][636][13]],
                        ['Harari',fed_data['DATA'][484][13]],
                        ['Konkani',fed_data['DATA'][563][13]],
                        ['Igbo',fed_data['DATA'][597][13]],
                        ['Rundi(Kirundi)',fed_data['DATA'][599][13]],
                        ['Shona',fed_data['DATA'][601][13]],
                        ['Pampangan(KapampanganPampango)',fed_data['DATA'][502][13]],
                        ['Kinyarwanda(Rwanda)',fed_data['DATA'][600][13]],
                        ['Swedish',fed_data['DATA'][552][13]],
                        ['Yiddish',fed_data['DATA'][554][13]],
                        ['HaitianCreole',fed_data['DATA'][508][13]],
                        ['Oriya(Odia)',fed_data['DATA'][566][13]],
                        ['Edo',fed_data['DATA'][592][13]],
                        ['Uyghur',fed_data['DATA'][635][13]],
                        ['Danish',fed_data['DATA'][546][13]],
                        ['MinDong',fed_data['DATA'][618][13]],
                        ['Kabyle',fed_data['DATA'][472][13]],
                        ['Mongolian',fed_data['DATA'][588][13]],
                        ['Lingala',fed_data['DATA'][598][13]],
                        ['Ganda',fed_data['DATA'][596][13]],
                        ['Mohawk',fed_data['DATA'][438][13]],
                        ['Belarusan',fed_data['DATA'][526][13]],
                        ['Ga',fed_data['DATA'][595][13]],
                        ['Bilen',fed_data['DATA'][475][13]],
                        ['Vlaams(Flemish)',fed_data['DATA'][553][13]],
                        ['Inuktitut',fed_data['DATA'][434][13]],
                        ['Catalan',fed_data['DATA'][577][13]],
                        ['Ewe',fed_data['DATA'][593][13]],
                        ['Wolof',fed_data['DATA'][603][13]],
                        ['Dinka',fed_data['DATA'][607][13]],
                        ['Fulah(PularPulaarFulfulde)',fed_data['DATA'][594][13]],
                        ['SwampyCree',fed_data['DATA'][397][13]],
                        ['Bikol',fed_data['DATA'][495][13]],
                        ['Pangasinan',fed_data['DATA'][503][13]],
                        ['Waray-Waray',fed_data['DATA'][505][13]],
                        ['Frisian',fed_data['DATA'][548][13]],
                        ['Ottawa(Odawa)',fed_data['DATA'][407][13]],
                        ['Norwegian',fed_data['DATA'][551][13]],
                        ['Kashmiri',fed_data['DATA'][562][13]],
                        ['ScottishGaelic',fed_data['DATA'][541][13]],
                        ['Icelandic',fed_data['DATA'][550][13]],
                        ['Malagasy',fed_data['DATA'][500][13]],
                        ['Welsh',fed_data['DATA'][542][13]],
                        ['Bamanankan',fed_data['DATA'][591][13]],
                        ['QuebecSignLanguage',fed_data['DATA'][611][13]],
                        ['NorthernEastCree',fed_data['DATA'][394][13]],
                        ['Oneida',fed_data['DATA'][439][13]],
                        ['Fijian',fed_data['DATA'][497][13]],
                        ['MooseCree',fed_data['DATA'][392][13]],
                        ['Naskapi',fed_data['DATA'][393][13]],
                        ['Montagnais(Innu)',fed_data['DATA'][391][13]],
                        ['WoodsCree',fed_data['DATA'][398][13]],
                        ['MikmaqCHANGESPELLING',fed_data['DATA'][402][13]],
                        ['Algonquin',fed_data['DATA'][404][13]],
                        ['Dene',fed_data['DATA'][415][13]],
                        ['GwichCHANGEPELLING',fed_data['DATA'][417][13]],
                        ['Sarsi(Sarcee)',fed_data['DATA'][418][13]],
                        ['Cayuga',fed_data['DATA'][437][13]],
                        ['Michif',fed_data['DATA'][442][13]],
                        ['Dakota',fed_data['DATA'][454][13]],
                        ['Nuu-chah-nulth(Nootka)',fed_data['DATA'][466][13]],
                        ['Blackfoot',fed_data['DATA'][388][13]],
                        ['Atikamekw',fed_data['DATA'][390][13]],
                        ['PlainsCree',fed_data['DATA'][395][13]],
                        ['SouthernEastCree',fed_data['DATA'][396][13]],
                        ['Malecite',fed_data['DATA'][401][13]],
                        ['Babine(WetsuwetenCHANGESPELLING)',fed_data['DATA'][411][13]],
                        ['Beaver',fed_data['DATA'][412][13]],
                        ['Carrier',fed_data['DATA'][413][13]],
                        ['Chilcotin',fed_data['DATA'][414][13]],
                        ['Dogrib(Tlicho)',fed_data['DATA'][416][13]],
                        ['Sekani',fed_data['DATA'][419][13]],
                        ['NorthSlavey(Hare)',fed_data['DATA'][421][13]],
                        ['SouthSlavey',fed_data['DATA'][422][13]],
                        ['Slaveyn.o.s.',fed_data['DATA'][423][13]],
                        ['Kaska(Nahani)',fed_data['DATA'][425][13]],
                        ['Tahltan',fed_data['DATA'][426][13]],
                        ['NorthernTutchone',fed_data['DATA'][428][13]],
                        ['SouthernTutchone',fed_data['DATA'][429][13]],
                        ['Haida',fed_data['DATA'][431][13]],
                        ['Inuinnaqtun(Inuvialuktun)',fed_data['DATA'][433][13]],
                        ['Kutenai',fed_data['DATA'][441][13]],
                        ['Comox',fed_data['DATA'][444][13]],
                        ['Halkomelem',fed_data['DATA'][445][13]],
                        ['Lillooet',fed_data['DATA'][446][13]],
                        ['Okanagan',fed_data['DATA'][447][13]],
                        ['Shuswap(Secwepemctsin)',fed_data['DATA'][448][13]],
                        ['Squamish',fed_data['DATA'][449][13]],
                        ['Straits',fed_data['DATA'][450][13]],
                        ['Thompson(Ntlakapamux)',fed_data['DATA'][451][13]],
                        ['Stoney',fed_data['DATA'][455][13]],
                        ['Tlingit',fed_data['DATA'][457][13]],
                        ['Gitxsan(Gitksan)',fed_data['DATA'][459][13]],
                        ['NisgaaCHANGESPELLING',fed_data['DATA'][460][13]],
                        ['Tsimshian',fed_data['DATA'][461][13]],
                        ['Haisla',fed_data['DATA'][463][13]],
                        ['Heiltsuk',fed_data['DATA'][464][13]],
                        ['Kwakiutl(KwakwalaCHANGESPELLING)',fed_data['DATA'][465][13]]]
        language_groups.sort(key= lambda x:x[1], reverse=True)
        seniors_count = fed_data['DATA'][36][13]
        median_age = fed_data['DATA'][39][13]

        # Get Prov 2016 Census Data
        prov_id = '2016A000235'
        with urllib.request.urlopen('https://www12.statcan.gc.ca/rest/census-recensement/CPR2016.json?lang=E&dguid='+prov_id) as url:
                prov_data = json.loads(url.read().decode())
        total_pop_count_prov = prov_data['DATA'][0][13]
        immigrants_count_prov = prov_data['DATA'][1141][13]
        immigrants_2011to2016_count_prov = prov_data['DATA'][1148][13]
        non_citizen_count_prov = prov_data['DATA'][1138][13]
        industry_total_prov =  prov_data['DATA'][1898][13]
        industry_groups_prov = {
                        'agriculture, forestry, fishing and hunting':prov_data['DATA'][1899][13],
                        'mining, quarrying, and oil and gas extraction':prov_data['DATA'][1900][13],
                        'utilities':prov_data['DATA'][1901][13],
                        'construction':prov_data['DATA'][1902][13],
                        'manufacturing':prov_data['DATA'][1903][13],
                        'wholesale trade':prov_data['DATA'][1904][13],
                        'retail trade':prov_data['DATA'][1905][13],
                        'transportation and warehousing':prov_data['DATA'][1906][13],
                        'information and cultural industries':prov_data['DATA'][1907][13],
                        'finance and insurance':prov_data['DATA'][1908][13],
                        'real estate and rental and leasing':prov_data['DATA'][1909][13],
                        'professional, scientific and technical services':prov_data['DATA'][1910][13],
                        'management of companies and enterprises':prov_data['DATA'][1911][13],
                        'administrative and support, waste management and remediation services':prov_data['DATA'][1912][13],
                        'educational services':prov_data['DATA'][1913][13],
                        'health care and social assistance':prov_data['DATA'][1914][13],
                        'arts, entertainment and recreation':prov_data['DATA'][1915][13],
                        'accommodation and food services':prov_data['DATA'][1916][13],
                        'other services (except public administration)':prov_data['DATA'][1917][13],
                        'public administration':prov_data['DATA'][1918][13]
                        }

        # Get FED 2021 Census Data
        # Could possibly remove this part - Ramsha 
        with urllib.request.urlopen('https://api.statcan.gc.ca/census-recensement/profile/sdmx/rest/data/STC_CP,DF_FED/A5.'+fed_id+
                                '.1.1.1?format=jsondata&detail=dataonly') as url:
                fed_data_2021 = json.loads(url.read().decode())
        total_pop_count_2021 = float(fed_data_2021['data']['dataSets'][0]['series']['0:0:0:0:0']['observations']['0'][0])
        with urllib.request.urlopen('https://api.statcan.gc.ca/census-recensement/profile/sdmx/rest/data/STC_CP,DF_FED/A5.'+fed_id+
                                '.1.3.1?format=jsondata&detail=dataonly') as url:
                fed_data_2021 = json.loads(url.read().decode())
        total_pop_growth_rate_2016to2021 = float(fed_data_2021['data']['dataSets'][0]['series']['0:0:0:0:0']['observations']['0'][0])
        
        #Get FED 2021 Census Data for age from Table: 98-10-0025-01
        
        df = StatcanTable(98100025)
        df.columns = ['REF_DATE','GEO','DGUID','AGE','COORDINATE','GENDER-TOTAL','SYMBOL','GENDER-MEN+','SYMBOL1','GENDER-WOMEN+','SYMBOL2']
        
        on_med_age_21 = df.loc[(df['DGUID']=='2021A000235') & (df['AGE']=="Median age"),'GENDER-TOTAL'].iloc[0]
        on_seniors_pop_21 = df.loc[(df['DGUID']=='2021A000235') & (df['AGE']=="65 years and over"),'GENDER-TOTAL'].iloc[0]
        on_total_pop_21 = df.loc[(df['DGUID']=='2021A000235') & (df['AGE']=="Total - Age"),'GENDER-TOTAL'].iloc[0]
        on_seniors_prop_21 = round(on_seniors_pop_21/on_total_pop_21 * 100 ,1)
            
        fed_med_age_21 = df.loc[(df['DGUID']==fed_id) & (df['AGE']=="Median age"),'GENDER-TOTAL'].iloc[0]
        seniors_pop_21 = df.loc[(df['DGUID']==fed_id) & (df['AGE']=="65 years and over"),'GENDER-TOTAL'].iloc[0]
        total_pop_21 = df.loc[(df['DGUID']==fed_id) & (df['AGE']=="Total - Age"),'GENDER-TOTAL'].iloc[0]
        seniors_prop_21 = round(seniors_pop_21/total_pop_21 * 100,1)


        #### INCOME FOR FEDs, zach ####

        # 1) Go to https://www12.statcan.gc.ca/census-recensement/2021/dp-pd/prof/details/download-telecharger.cfm?Lang=E #
        # 2) Download Canada, provinces, territories and federal electoral districts (FEDs) (2013 Representation Order), CSV #
        # 3) Place in same folder as .py file #
        # 4) Change first line below to directory you are using #
  
        Census21 = pd.read_csv (r"C:\Users\ramsha.jaweed\Documents\98-401-X2021010_English_CSV_data.csv",low_memory=False, encoding='latin-1')
            
        on_med_tot_inc = Census21.loc[(Census21['CHARACTERISTIC_NAME'].str.contains("Median total income of household in 2020")) & (Census21['GEO_NAME']=='Ontario'),'C1_COUNT_TOTAL'].iloc[0]
        on_prevalence_lim = Census21.loc[(Census21['CHARACTERISTIC_NAME'].str.contains("Prevalence of low income based on the Low-income measure, after tax")) & (Census21['GEO_NAME']=='Ontario'),'C1_COUNT_TOTAL'].iloc[0]

        Census_fed = Census21.loc[Census21['DGUID']==fed_id]
        fed_med_tot_inc = Census_fed.loc[Census_fed['CHARACTERISTIC_NAME'].str.contains("Median total income of household in 2020"),'C1_COUNT_TOTAL'].iloc[0]
        fed_prevalence_lim = Census_fed.loc[Census_fed['CHARACTERISTIC_NAME'].str.contains("Prevalence of low income based on the Low-income measure, after tax"),'C1_COUNT_TOTAL'].iloc[0]

      
        # Generate DOCX Document
        document = Document()
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
        
        ############################################################################↓↓↓↓↓FORMATTING
        style = document.styles["Normal"]
        font = style.font
        font.name = 'Arial'

        style = document.styles['Heading 4']
        font = style.font
        rFonts = font.element.rPr.rFonts
        rFonts.set(qn("w:asciiTheme"), "Arial")
        font.name = 'Arial'
        font.size = Pt(28)
        font.color.rgb = RGBColor(192,0,0)
        font.italic = False
        font.bold = False

        style = document.styles['Title']
        font = style.font
        rFonts = font.element.rPr.rFonts
        rFonts.set(qn("w:asciiTheme"), "Arial")
        font.name = 'Arial'
        font.size = Pt(14)
        font.color.rgb = RGBColor(192,0,0)
        font.italic = False
        font.bold = True

        style = document.styles['Heading 3']
        font = style.font
        rFonts = font.element.rPr.rFonts
        rFonts.set(qn("w:asciiTheme"), "Arial")
        font.name = 'Arial'
        font.size = Pt(11)
        font.color.rgb = RGBColor(192,0,0)
        font.italic = False
        font.bold = True

        par1=document.add_paragraph(f'Profile:',style="Heading 4")
        par1.style = document.styles['Heading 4']

        para=document.add_paragraph("Constituency:",style="Title")
        ############################################################################↑↑↑↑↑FORMATTING
        
        document.add_paragraph('Political Affiliation:',style="Heading 3")
        document.add_paragraph('Parliamentary Tenure:',style="Heading 3")
        document.add_paragraph('Notable Parliametary Roles:',style="Heading 3")
        document.add_paragraph('Background:',style="Heading 3")
        document.add_paragraph('Media Scan:',style="Heading 3")
        p=document.add_paragraph('Constituency Demographics:',style="Heading 3").add_run()
        p.add_break()
        p=document.add_paragraph().add_run('Geography')
        p.bold=True
        p=document.add_paragraph().add_run('Population')
        p.bold=True
        document.add_paragraph(
                        f'According to the 2021 Census, the population of the {fed_name} FED was {total_pop_count_2021:,.0f}, '+
                        ('an increase ' if total_pop_growth_rate_2016to2021>0 else 'a decrease ')+
                        f'of {total_pop_growth_rate_2016to2021/100:.1%} from 2016.  '+
                        ('This is in contrast to the province where the population grew by 775,448 (5.8%).' if total_pop_growth_rate_2016to2021 != 5.8 else
                        'The population of the province grew by 775,448, which was also an increase of 5.8%.'))

        p=document.add_paragraph().add_run('Indigenous')
        p.bold=True
        document.add_paragraph(
                        f'In 2016, {indigenous_count:,.0f} individuals in the {fed_name} FED self-identified as Indigenous, '+
                        f'representing about {indigenous_count/total_pop_count:.1%} '+
                        f'of the population. In comparison, 2.8% (374,395) of the provincial population self-identified as Indigenous.')

        p=document.add_paragraph().add_run('Visible Minorities')
        p.bold=True
        document.add_paragraph(
                        f'Roughly {vis_minority_count/vis_minority_total:.1%} of the {fed_name} FED population self-identified as a '+
                        f'visible minority in 2016, or {vis_minority_count:,.0f}  '+
                        f'individuals. '+
                        f'{vis_minority_groups[0][0]} ({vis_minority_groups[0][1]:,.0f}; {vis_minority_groups[0][1]/vis_minority_total:.1%}), '+
                        f'{vis_minority_groups[1][0]} ({vis_minority_groups[1][1]:,.0f}; {vis_minority_groups[1][1]/vis_minority_total:.1%}), '+
                        f'and '+
                        f'{vis_minority_groups[2][0]} ({vis_minority_groups[2][1]:,.0f}; {vis_minority_groups[2][1]/vis_minority_total:.1%}) '+
                        f'were the top three ethnicities. About 29.3% or 3,885,585 individuals in Ontario identified  '+
                        f'as a visible minority, and the top three ethnicities were South Asian (1,150,415; 8.7%), Chinese (754,550; 5.7%) and Black (627,715; 4.7%).')

        p=document.add_paragraph().add_run('Persons with Disabilities')
        p.bold=True
        document.add_paragraph(
                        f'While data for the {fed_name} FED is scarce, ____ (____%) of the ____ CMA population identified as persons '+
                        f'with disabilities in 2017. Comparatively, 2,616,170 '+
                        f'(24.1%) individuals identified as persons with disabilities at the provincial level.')

        p=document.add_paragraph().add_run('Newcomers')
        p.bold=True
        document.add_paragraph(
                        f'About {immigrants_count/immigrants_total:.1%} ({immigrants_count:,.0f}) of the {fed_name} FED’s population were '+
                        f'immigrants at the time of the 2016 Census. About '+
                        f'{immigrants_2011to2016_count/immigrants_total:.1%} ({immigrants_2011to2016_count:,.0f}) immigrated between '+
                        f'2011 and 2016. In comparison, 29.1% (3,852,145) of Ontario’s '+
                        f'population were immigrants at the time of the Census and 3.6% (472,170) immigrated into the province between '+
                        f'2011 and 2016. Roughly '+
                        f'{non_citizen_count/immigrants_total:.1%} ({non_citizen_count:,.0f}) of the {fed_name} FED’s population and '+
                        f'7.7% (1,019,095) of Ontario’s population did not have a '+
                        f'Canadian Citizenship at the time of the 2016 Census.')

        p=document.add_paragraph().add_run('Languages')
        p.bold=True
        document.add_paragraph(
                        f'Approximately {englishandfrench/language_total:.1%} ({englishandfrench:,.0f}) of the {fed_name} FED '+
                        f'population had knowledge of both official languages at the time '+
                        f'of the 2016 Census, compared to 11.2% (1,490,390) of the province. Optional sentence: ____ is an official '+
                        f'language minority community with at least '+
                        f'one school where the primary teaching language is the official language minority language. '+
                        f'Apart from official languages, the top three languages most '+
                        f'spoken at home in the FED were '+
                        f'{language_groups[0][0]} ({language_groups[0][1]:,.0f}; {language_groups[0][1]/language_total:.1%}), '+
                        f'{language_groups[1][0]} ({language_groups[1][1]:,.0f}; {language_groups[1][1]/language_total:.1%}), and '+
                        f'{language_groups[2][0]} ({language_groups[2][1]:,.0f}; {language_groups[2][1]/language_total:.1%}) '+
                        f'Comparatively, the top three languages most spoken at home in Ontario were '+
                        f'Mandarin (220,535; 1.7%), Cantonese (198,745; 1.5%), and Punjabi (132,135; 1%).')

        # VARIABLES FOR SENIORS SECTION
        seniors_prop_21 = round(seniors_pop_21/total_pop_21 * 100,1)
        on_seniors_prop_21 = round(on_seniors_pop_21/on_total_pop_21 * 100 ,1)

        if fed_med_age_21 > on_med_age_21: #make sure on_med_age_21 is equal to 41.6
            choice_1='an older '
        elif fed_med_age_21 < on_med_age_21:
            choice_1='a younger '
        elif fed_med_age_21 == on_med_age_21:
            choice_1='a similar '

        if seniors_prop_21<(on_seniors_prop_21-0.2):
            choice_2=f'There were proportionally fewer seniors, with {seniors_prop_21}% of the population 65 years or older, versus {on_seniors_prop_21}% for the province.'
        elif (on_seniors_prop_21-0.2)<=seniors_prop_21<=(on_seniors_prop_21+0.2):
            choice_2=f"The proportion of seniors was similar to that of the province. About {seniors_prop_21}% of the {fed_name} FED's population was 65 years or older, versus {on_seniors_prop_21}% for all of Ontario."
        elif seniors_prop_21>(on_seniors_prop_21+0.2):
            choice_2=f'There were proportionally more seniors, with {seniors_prop_21}% of the population 65 years or older, versus {on_seniors_prop_21}% for the province.'
            
        #CREATING SENIORS SECTION PARAGRAPH
        p=document.add_paragraph().add_run('Seniors')
        p.bold=True
        document.add_paragraph(
                        f'The {fed_name} FED had '+
                        str(choice_1)+
                        f'age profile compared to the province. According to the 2021 Census, the median age was '+
                        str(fed_med_age_21)+
                        f', compared with {on_med_age_21} for Ontario. '+
                        choice_2)
        #END SENIORS SECTION
        
        p=document.add_paragraph().add_run('Labour Force')
        p.bold=True
        document.add_paragraph(
                        f'According to the Labour Force Survey, the ____ CMA had a labour force of about ____ in February 2022. '+
                        f'The unemployment rate was ____% — ____ '+
                        f'than Ontario’s 5.5%. The participation rate and the employment rate in the CMA were ____ the province, '+
                        f'at ____% and ____% respectively. By comparison, Ontario’s participation rate was 65.4%, '+
                        f'and the employment rate was 61.7%. '+
                        f'Labour market conditions in the ____ CMA ____ between February 2021 and February 2022 as employment '+
                        f'____ by ____. The unemployment rate in the ____ CMA ____ by ____ percentage points to ____% over the year, '+
                        f'compared to a decrease of 3.7 percentage points provincially.')

        #### Income Paragraph for FED, zach ####
        p=document.add_paragraph().add_run('Income')
        p.bold=True
        document.add_paragraph(
                        f'The median total income for the {fed_name} FED private households in 2020 was ${fed_med_tot_inc:,.0f} which was '+
                        ('higher ' if  (fed_med_tot_inc>on_med_tot_inc) else 'lower ')+
                        f'than Ontario ($91,000). '+
                        f'The prevalence of low-income based on the Low-income measure after-tax (LIM-AT) indicator is {fed_prevalence_lim}% '+
                        f'in the {fed_name} FED compared to 10.1% for Ontario.')  
        
        p=document.add_paragraph().add_run('Industrial Base')
        p.bold=True
        over_rep_industries = []
        for x in industry_groups:
                x_industry_name = x[0]
                x_prov_pct = industry_groups_prov[x_industry_name]/industry_total_prov
                x_fed_pct = x[1]/industry_total
                if x_fed_pct/x_prov_pct>=2:
                        over_rep_industries.append([x_industry_name,x_prov_pct,x_fed_pct])
        over_rep_industries.sort(key= lambda x:x[2]/x[1], reverse=True)
        document.add_paragraph(
                        f'According to the 2016 Census, the largest three sectors by employment size in the {fed_name} FED were '+
                        f'{industry_groups[0][0]} ({industry_groups[0][1]:,.0f}; {industry_groups[0][1]/industry_total:.1%}), '+
                        f'{industry_groups[1][0]} ({industry_groups[1][1]:,.0f}; {industry_groups[1][1]/industry_total:.1%}), and '+
                        f'{industry_groups[2][0]} ({industry_groups[2][1]:,.0f}; {industry_groups[2][1]/industry_total:.1%}). '+
                        f'Employment in these top three sectors made up about '+
                        f'{industry_groups[0][1]/industry_total+industry_groups[1][1]/industry_total+industry_groups[2][1]/industry_total:.1%} '+
                        f'of the employed labour force in the {fed_name} FED, '+
                        ('higher' if ((industry_groups[0][1]+industry_groups[1][1]+industry_groups[2][1])/industry_total>(industry_groups_prov[industry_groups[0][0]]+industry_groups_prov[industry_groups[1][0]]+industry_groups_prov[industry_groups[2][0]])/industry_total_prov)
                         else 'lower')+
                        f' than the Ontario share of '+
                        f'{industry_groups_prov[industry_groups[0][0]]/industry_total_prov+industry_groups_prov[industry_groups[1][0]]/industry_total_prov+industry_groups_prov[industry_groups[2][0]]/industry_total_prov:.1%}. '+
                        ('Additionally, ' if len(over_rep_industries)>0 else '')+ 
                        (', '.join([x[0] for x in over_rep_industries[:-1]])+', and '+over_rep_industries[-1][0] if len(over_rep_industries)>2 else '')+
                        (over_rep_industries[0][0]+' and '+over_rep_industries[1][0] if len(over_rep_industries)==2 else '')+
                        (over_rep_industries[0][0] if len(over_rep_industries)==1 else '')+
                        (' is ' if len(over_rep_industries)==1 else ' are ' if len(over_rep_industries)>1 else '')+
                        (f'overrepresented in the {fed_name} FED. ' if len(over_rep_industries)>0 else '')+
                        (''.join([x[0].capitalize()+f' represents about {x[2]:.1%} of the employment in the FED compared to {x[1]:.1%} at the provincial level. ' for x in over_rep_industries])))

        document.save(fed_name+'.docx')

        


# In[ ]:


## Produces a file for CMA level data       
  
# Generate DOCX Document
document = Document()
sections = document.sections
for section in sections:
  section.top_margin = Cm(1.27)
  section.bottom_margin = Cm(1.27)
  section.left_margin = Cm(1.27)
  section.right_margin = Cm(1.27)
      
############################################################################↓↓↓↓↓FORMATTING
style = document.styles["Normal"]
font = style.font
font.name = 'Arial'

style = document.styles['Heading 4']
font = style.font
rFonts = font.element.rPr.rFonts
rFonts.set(qn("w:asciiTheme"), "Arial")
font.name = 'Arial'
font.size = Pt(28)
font.color.rgb = RGBColor(192,0,0)
font.italic = False
font.bold = False

style = document.styles['Title']
font = style.font
rFonts = font.element.rPr.rFonts
rFonts.set(qn("w:asciiTheme"), "Arial")
font.name = 'Arial'
font.size = Pt(14)
font.color.rgb = RGBColor(192,0,0)
font.italic = False
font.bold = True

style = document.styles['Heading 3']
font = style.font
rFonts = font.element.rPr.rFonts
rFonts.set(qn("w:asciiTheme"), "Arial")
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = RGBColor(192,0,0)
font.italic = False
font.bold = True

par1=document.add_paragraph(f'Labour Force Statistics:',style="Heading 4")
par1.style = document.styles['Heading 4']

cma_id_list_2016 = ['2016S0503568' , '2016S0503522', '2016S0503543', '2016S0503580', '2016S0503550', '2016S0503537', 
                  '2016S0503521', '2016S0503541', '2016S0503555', '2016S0503532', '2016S050535505', '2016S0503529',
                  '2016S0503559', '2016S0503595', '2016S0503535', '2016S0503559']

      #####################################################################
      ######################LFS Data for the CMA###########################
      #####################################################################
for i in range(0,16):
      cma_id = cma_id_list_2016[i]        
          
      lfs_cma = StatcanTable(14100380)
  
      #Filtering data for ON CMAs only, estimates and seasonally adjusted 
      lfs_cma = lfs_cma[lfs_cma['GEO'].str.contains("Ontario")]
      lfs_cma = lfs_cma.loc[(lfs_cma['Statistics']=='Estimate') & (lfs_cma['Data type']=='Seasonally adjusted')]
          
      #Filtering lfs_cma to get data only for selected CMA:

      lfs_cma = lfs_cma.loc[(lfs_cma['DGUID']==cma_id)]
      cma_name = lfs_cma.loc[lfs_cma['Labour force characteristics']=="Labour force",'GEO'].iloc[0][:-9]
      #changing the date format for REF_DATE column 
      lfs_cma['REF_DATE'] = pd.to_datetime(lfs_cma['REF_DATE'])
      ly = max(lfs_cma['REF_DATE']) - relativedelta(years=1)
      month = max(lfs_cma['REF_DATE'])
      #Filtering CMA data to get ON unemployment rate, participation rate and employment rate
      cma_labour_force = lfs_cma.loc[(lfs_cma['Labour force characteristics']=="Labour force") & (lfs_cma['REF_DATE']==max(lfs_cma['REF_DATE'])),'VALUE'].iloc[0]
      cma_unemp_rate = lfs_cma.loc[(lfs_cma['Labour force characteristics']=="Unemployment rate") & (lfs_cma['REF_DATE']==max(lfs_cma['REF_DATE'])),'VALUE'].iloc[0]
      cma_part_rate = lfs_cma.loc[(lfs_cma['Labour force characteristics']=="Participation rate") & (lfs_cma['REF_DATE']==max(lfs_cma['REF_DATE'])),'VALUE'].iloc[0]
      cma_emp_rate = lfs_cma.loc[(lfs_cma['Labour force characteristics']=="Employment rate") & (lfs_cma['REF_DATE']==max(lfs_cma['REF_DATE'])),'VALUE'].iloc[0]
      cma_unemp_rate_change = lfs_cma.loc[(lfs_cma['Labour force characteristics']=="Unemployment rate") & (lfs_cma['REF_DATE']==ly),'VALUE'].iloc[0]         - cma_unemp_rate
      cma_emp_change = lfs_cma.loc[(lfs_cma['Labour force characteristics']=="Employment") & (lfs_cma['REF_DATE']==max(lfs_cma['REF_DATE'])),'VALUE'].iloc[0]         - lfs_cma.loc[(lfs_cma['Labour force characteristics']=="Employment") & (lfs_cma['REF_DATE']==ly),'VALUE'].iloc[0]
      
      if i == 0:
          par1=document.add_paragraph(f'Month:{month}',style="Heading 4")
          par1.style = document.styles['Heading 4']
          pass
          
      
      # Define variables for Ontario's rates 
      lfs_on = StatcanTable(14100287)
      lfs_on = lfs_on.loc[(lfs_on['Statistics']=='Estimate') & (lfs_on['Data type']=='Seasonally adjusted')                           & (lfs_on['GEO']=='Ontario') & (lfs_on['Sex']=='Both sexes') & (lfs_on['Age group']=='15 years and over')]
      #changing the date format for REF_DATE column
      lfs_on['REF_DATE'] = pd.to_datetime(lfs_on['REF_DATE'])
      ly = max(lfs_on['REF_DATE']) - relativedelta(years=1)
      
      on_unemp_rate = lfs_on.loc[(lfs_on['Labour force characteristics']=="Unemployment rate") & (lfs_on['REF_DATE']==max(lfs_on['REF_DATE'])),'VALUE'].iloc[0]
      on_part_rate = lfs_on.loc[(lfs_on['Labour force characteristics']=="Participation rate") & (lfs_on['REF_DATE']==max(lfs_on['REF_DATE'])),'VALUE'].iloc[0]
      on_emp_rate = lfs_on.loc[(lfs_on['Labour force characteristics']=="Employment rate") & (lfs_on['REF_DATE']==max(lfs_on['REF_DATE'])),'VALUE'].iloc[0]
      on_unemp_rate_change = lfs_on.loc[(lfs_on['Labour force characteristics']=="Unemployment rate") & (lfs_on['REF_DATE']==ly),'VALUE'].iloc[0]         - on_unemp_rate
      

      p=document.add_paragraph().add_run( f'{cma_name} Labour Force Characteristics')
      p.bold=True
      document.add_paragraph(
                      f'According to the Labour Force Survey, the {cma_name} CMA had a labour force of about {(cma_labour_force *1000):,.0f} in {month} 2022. '+
                      f'The unemployment rate was {cma_unemp_rate}% —'+ ('higher' if (cma_unemp_rate>on_unemp_rate) else 'lower') +
                      f' than Ontario’s {on_unemp_rate}%.' +
                      f'The participation rate and the employment rate in the CMA were ' + ('higher' if (cma_emp_rate>on_emp_rate) else 'lower') +
                      f' than the province, '+
                      f'at {cma_emp_rate}% and {cma_part_rate}% respectively. By comparison, Ontario’s participation rate was {on_part_rate}%, '+
                      f'and the employment rate was {on_emp_rate}%. '+
                      f'Labour market conditions in the {cma_name} CMA ' + ('improved' if (cma_emp_change>0) else 'declined') +
                      f' between {month} 2021 and {month} 2022 as employment '+
                      ('grew' if (cma_emp_change >0) else 'shrank') +
                      f' by {abs(cma_emp_change*1000):,.0f}. The unemployment rate in the {cma_name} CMA ' +
                      ('increased' if (cma_unemp_rate_change>0) else 'decreased') +
                      f' by {abs(cma_unemp_rate_change):,.1f} percentage points to {cma_unemp_rate}% over the year, '+
                      f'compared to a' + ('n increase ' if (on_unemp_rate_change>0) else 'decrease') +
                      f'of {abs(on_unemp_rate_change):,.1f} percentage points provincially.')
document.save('LFS_CMA' + '.docx')

