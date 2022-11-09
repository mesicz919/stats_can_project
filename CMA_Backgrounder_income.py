#!/usr/bin/env python
# coding: utf-8

# In[1]:


import urllib.request, json
from docx import Document

import requests
import pandas as pd
import io
import json
import os
import zipfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta



# In[2]:


#define function for API to download statcan tables
def StatcanTable(TableNumber):
    TableNumber = str(TableNumber)
    r = requests.get("https://www150.statcan.gc.ca/n1/en/tbl/csv/" + TableNumber + "-eng.zip", stream=True)
    #print(r)
    z = zipfile.ZipFile(io.BytesIO(r.content))
    z.extractall(os.getcwd())
    RawData = pd.read_csv(TableNumber + ".csv",low_memory=False)
    return RawData


# In[6]:


############################################################################
# Generate Doc1 using 2016 census data and 2016 DGUIDs for CMAs
# Sections: Indigenous, Minorities, Newcomers, Languages, Industrial Base
############################################################################
cma_id_list_2016 = ['2016S0503568' , '2016S0503522', '2016S0503543', '2016S0503580', '2016S0503550', '2016S0503537', 
                    '2016S0503521', '2016S0503541', '2016S0503555', '2016S0503532', '2016S050535505', '2016S0503529',
                    '2016S0503559', '2016S0503595', '2016S0503535', '2016S0503559']

cma_id_list_2021 = ['2021S0503568','2021S0503522', '2021S0503543', '2021S0503580', '2021S0503550', '2021S0503537', '2021S0503521',
                    '2021S0503541','2021S0503555', '2021S0503532', '2021S0503505', '2021S0503529', '2021S0503539', '2021S0503595',
                    '2021S0503535', '2021S0503559']
 
for i in range(0,17):
        cma_id = cma_id_list_2016[i] 
    
        # Get cma 2016 Census Data
        with urllib.request.urlopen('https://www12.statcan.gc.ca/rest/census-recensement/CPR2016.json?lang=E&dguid='+cma_id) as url:
                cma_data = json.loads(url.read().decode())
        cma_name = cma_data['DATA'][0][4]
        total_pop_count = cma_data['DATA'][0][13]
        immigrants_total = cma_data['DATA'][1139][13]
        immigrants_count = cma_data['DATA'][1141][13]
        immigrants_2011to2016_count = cma_data['DATA'][1148][13]
        non_citizen_count = cma_data['DATA'][1138][13]
        indigenous_count = cma_data['DATA'][1289][13]
        vis_minority_count = cma_data['DATA'][1323][13]
        vis_minority_total =  cma_data['DATA'][1322][13]
        vis_minority_groups = [
                        ['South Asian',cma_data['DATA'][1324][13]],
                        ['Chinese',cma_data['DATA'][1325][13]],
                        ['Black',cma_data['DATA'][1326][13]],
                        ['Filipino',cma_data['DATA'][1327][13]],
                        ['Latin American',cma_data['DATA'][1328][13]],
                        ['Arab',cma_data['DATA'][1329][13]],
                        ['Southeast Asian',cma_data['DATA'][1330][13]],
                        ['West Asian',cma_data['DATA'][1331][13]],
                        ['Korean',cma_data['DATA'][1332][13]],
                        ['Japanese',cma_data['DATA'][1333][13]],
                        ['Multiracial',cma_data['DATA'][1335][13]]
                        ]
        vis_minority_groups.sort(key= lambda x:x[1], reverse=True)
        industry_total =  cma_data['DATA'][1898][13]
        industry_groups = [
                        ['agriculture, forestry, fishing and hunting',cma_data['DATA'][1899][13]],
                        ['mining, quarrying, and oil and gas extraction',cma_data['DATA'][1900][13]],
                        ['utilities',cma_data['DATA'][1901][13]],
                        ['construction',cma_data['DATA'][1902][13]],
                        ['manufacturing',cma_data['DATA'][1903][13]],
                        ['wholesale trade',cma_data['DATA'][1904][13]],
                        ['retail trade',cma_data['DATA'][1905][13]],
                        ['transportation and warehousing',cma_data['DATA'][1906][13]],
                        ['information and cultural industries',cma_data['DATA'][1907][13]],
                        ['finance and insurance',cma_data['DATA'][1908][13]],
                        ['real estate and rental and leasing',cma_data['DATA'][1909][13]],
                        ['professional, scientific and technical services',cma_data['DATA'][1910][13]],
                        ['management of companies and enterprises',cma_data['DATA'][1911][13]],
                        ['administrative and support, waste management and remediation services',cma_data['DATA'][1912][13]],
                        ['educational services',cma_data['DATA'][1913][13]],
                        ['health care and social assistance',cma_data['DATA'][1914][13]],
                        ['arts, entertainment and recreation',cma_data['DATA'][1915][13]],
                        ['accommodation and food services',cma_data['DATA'][1916][13]],
                        ['other services (except public administration)',cma_data['DATA'][1917][13]],
                        ['public administration',cma_data['DATA'][1918][13]]
                        ]
        industry_groups.sort(key= lambda x:x[1], reverse=True)
        englishandfrench = cma_data['DATA'][102][13]        
        language_total = cma_data['DATA'][99][13]    
        language_groups = [
                        ['Mandarin',cma_data['DATA'][617][13]],
                        ['Cantonese',cma_data['DATA'][615][13]],
                        ['Punjabi (Panjabi)',cma_data['DATA'][567][13]],
                        ['Spanish',cma_data['DATA'][581][13]],
                        ['Arabic',cma_data['DATA'][481][13]],
                        ['Urdu',cma_data['DATA'][570][13]],
                        ['Tamil',cma_data['DATA'][514][13]],
                        ['Persian(Farsi)',cma_data['DATA'][574][13]],
                        ['Tagalog(PilipinoFilipino)',cma_data['DATA'][504][13]],
                        ['Italian',cma_data['DATA'][578][13]],
                        ['Portuguese',cma_data['DATA'][579][13]],
                        ['Russian',cma_data['DATA'][533][13]],
                        ['Polish',cma_data['DATA'][532][13]],
                        ['Korean',cma_data['DATA'][586][13]],
                        ['Vietnamese',cma_data['DATA'][492][13]],
                        ['EnglishandFrench',cma_data['DATA'][645][13]],
                        ['Gujarati',cma_data['DATA'][560][13]],
                        ['German',cma_data['DATA'][549][13]],
                        ['Bengali',cma_data['DATA'][559][13]],
                        ['Hindi',cma_data['DATA'][561][13]],
                        ['Serbian',cma_data['DATA'][534][13]],
                        ['Greek',cma_data['DATA'][556][13]],
                        ['Romanian',cma_data['DATA'][580][13]],
                        ['Ukrainian',cma_data['DATA'][538][13]],
                        ['Somali',cma_data['DATA'][477][13]],
                        ['Hungarian',cma_data['DATA'][641][13]],
                        ['Croatian',cma_data['DATA'][529][13]],
                        ['Turkish',cma_data['DATA'][634][13]],
                        ['Chinesen.o.s.',cma_data['DATA'][621][13]],
                        ['AssyrianNeo-Aramaic',cma_data['DATA'][482][13]],
                        ['Albanian',cma_data['DATA'][519][13]],
                        ['Malayalam',cma_data['DATA'][513][13]],
                        ['Armenian',cma_data['DATA'][520][13]],
                        ['Pashto',cma_data['DATA'][573][13]],
                        ['Frenchandnon-officiallanguage',cma_data['DATA'][647][13]],
                        ['Macedonian',cma_data['DATA'][531][13]],
                        ['Japanese',cma_data['DATA'][583][13]],
                        ['MinNan(ChaochowTeochowFukienTaiwanese)',cma_data['DATA'][619][13]],
                        ['Telugu',cma_data['DATA'][515][13]],
                        ['Nepali',cma_data['DATA'][565][13]],
                        ['Amharic',cma_data['DATA'][480][13]],
                        ['Hebrew',cma_data['DATA'][485][13]],
                        ['Bulgarian',cma_data['DATA'][528][13]],
                        ['Sinhala(Sinhalese)',cma_data['DATA'][569][13]],
                        ['Dutch',cma_data['DATA'][547][13]],
                        ['Akan(Twi)',cma_data['DATA'][590][13]],
                        ['Ilocano',cma_data['DATA'][499][13]],
                        ['Creolen.o.s.',cma_data['DATA'][509][13]],
                        ['Tigrigna',cma_data['DATA'][487][13]],
                        ['Tibetan',cma_data['DATA'][626][13]],
                        ['Khmer(Cambodian)',cma_data['DATA'][491][13]],
                        ['Wu(Shanghainese)',cma_data['DATA'][620][13]],
                        ['Kurdish',cma_data['DATA'][572][13]],
                        ['Slovak',cma_data['DATA'][536][13]],
                        ['Ojibway',cma_data['DATA'][405][13]],
                        ['Bosnian',cma_data['DATA'][527][13]],
                        ['Czech',cma_data['DATA'][530][13]],
                        ['ChaldeanNeo-Aramaic',cma_data['DATA'][483][13]],
                        ['Oji-Cree',cma_data['DATA'][406][13]],
                        ['Sindhi',cma_data['DATA'][568][13]],
                        ['Hakka',cma_data['DATA'][616][13]],
                        ['Lao',cma_data['DATA'][629][13]],
                        ['Marathi',cma_data['DATA'][564][13]],
                        ['Cebuano',cma_data['DATA'][496][13]],
                        ['Creen.o.s.',cma_data['DATA'][399][13]],
                        ['Swahili',cma_data['DATA'][602][13]],
                        ['Serbo-Croatian',cma_data['DATA'][535][13]],
                        ['Finnish',cma_data['DATA'][640][13]],
                        ['Malay',cma_data['DATA'][501][13]],
                        ['AmericanSignLanguage',cma_data['DATA'][610][13]],
                        ['Slovene(Slovenian)',cma_data['DATA'][537][13]],
                        ['Lithuanian',cma_data['DATA'][524][13]],
                        ['Thai',cma_data['DATA'][630][13]],
                        ['Yoruba',cma_data['DATA'][604][13]],
                        ['Maltese',cma_data['DATA'][486][13]],
                        ['Kannada',cma_data['DATA'][512][13]],
                        ['Burmese',cma_data['DATA'][624][13]],
                        ['Hiligaynon',cma_data['DATA'][498][13]],
                        ['Latvian',cma_data['DATA'][523][13]],
                        ['Oromo',cma_data['DATA'][476][13]],
                        ['Estonian',cma_data['DATA'][639][13]],
                        ['Georgian',cma_data['DATA'][585][13]],
                        ['Azerbaijani',cma_data['DATA'][633][13]],
                        ['Afrikaans',cma_data['DATA'][545][13]],
                        ['Uzbek',cma_data['DATA'][636][13]],
                        ['Harari',cma_data['DATA'][484][13]],
                        ['Konkani',cma_data['DATA'][563][13]],
                        ['Igbo',cma_data['DATA'][597][13]],
                        ['Rundi(Kirundi)',cma_data['DATA'][599][13]],
                        ['Shona',cma_data['DATA'][601][13]],
                        ['Pampangan(KapampanganPampango)',cma_data['DATA'][502][13]],
                        ['Kinyarwanda(Rwanda)',cma_data['DATA'][600][13]],
                        ['Swedish',cma_data['DATA'][552][13]],
                        ['Yiddish',cma_data['DATA'][554][13]],
                        ['HaitianCreole',cma_data['DATA'][508][13]],
                        ['Oriya(Odia)',cma_data['DATA'][566][13]],
                        ['Edo',cma_data['DATA'][592][13]],
                        ['Uyghur',cma_data['DATA'][635][13]],
                        ['Danish',cma_data['DATA'][546][13]],
                        ['MinDong',cma_data['DATA'][618][13]],
                        ['Kabyle',cma_data['DATA'][472][13]],
                        ['Mongolian',cma_data['DATA'][588][13]],
                        ['Lingala',cma_data['DATA'][598][13]],
                        ['Ganda',cma_data['DATA'][596][13]],
                        ['Mohawk',cma_data['DATA'][438][13]],
                        ['Belarusan',cma_data['DATA'][526][13]],
                        ['Ga',cma_data['DATA'][595][13]],
                        ['Bilen',cma_data['DATA'][475][13]],
                        ['Vlaams(Flemish)',cma_data['DATA'][553][13]],
                        ['Inuktitut',cma_data['DATA'][434][13]],
                        ['Catalan',cma_data['DATA'][577][13]],
                        ['Ewe',cma_data['DATA'][593][13]],
                        ['Wolof',cma_data['DATA'][603][13]],
                        ['Dinka',cma_data['DATA'][607][13]],
                        ['Fulah(PularPulaarFulfulde)',cma_data['DATA'][594][13]],
                        ['SwampyCree',cma_data['DATA'][397][13]],
                        ['Bikol',cma_data['DATA'][495][13]],
                        ['Pangasinan',cma_data['DATA'][503][13]],
                        ['Waray-Waray',cma_data['DATA'][505][13]],
                        ['Frisian',cma_data['DATA'][548][13]],
                        ['Ottawa(Odawa)',cma_data['DATA'][407][13]],
                        ['Norwegian',cma_data['DATA'][551][13]],
                        ['Kashmiri',cma_data['DATA'][562][13]],
                        ['ScottishGaelic',cma_data['DATA'][541][13]],
                        ['Icelandic',cma_data['DATA'][550][13]],
                        ['Malagasy',cma_data['DATA'][500][13]],
                        ['Welsh',cma_data['DATA'][542][13]],
                        ['Bamanankan',cma_data['DATA'][591][13]],
                        ['QuebecSignLanguage',cma_data['DATA'][611][13]],
                        ['NorthernEastCree',cma_data['DATA'][394][13]],
                        ['Oneida',cma_data['DATA'][439][13]],
                        ['Fijian',cma_data['DATA'][497][13]],
                        ['MooseCree',cma_data['DATA'][392][13]],
                        ['Naskapi',cma_data['DATA'][393][13]],
                        ['Montagnais(Innu)',cma_data['DATA'][391][13]],
                        ['WoodsCree',cma_data['DATA'][398][13]],
                        ['MikmaqCHANGESPELLING',cma_data['DATA'][402][13]],
                        ['Algonquin',cma_data['DATA'][404][13]],
                        ['Dene',cma_data['DATA'][415][13]],
                        ['GwichCHANGEPELLING',cma_data['DATA'][417][13]],
                        ['Sarsi(Sarcee)',cma_data['DATA'][418][13]],
                        ['Cayuga',cma_data['DATA'][437][13]],
                        ['Michif',cma_data['DATA'][442][13]],
                        ['Dakota',cma_data['DATA'][454][13]],
                        ['Nuu-chah-nulth(Nootka)',cma_data['DATA'][466][13]],
                        ['Blackfoot',cma_data['DATA'][388][13]],
                        ['Atikamekw',cma_data['DATA'][390][13]],
                        ['PlainsCree',cma_data['DATA'][395][13]],
                        ['SouthernEastCree',cma_data['DATA'][396][13]],
                        ['Malecite',cma_data['DATA'][401][13]],
                        ['Babine(WetsuwetenCHANGESPELLING)',cma_data['DATA'][411][13]],
                        ['Beaver',cma_data['DATA'][412][13]],
                        ['Carrier',cma_data['DATA'][413][13]],
                        ['Chilcotin',cma_data['DATA'][414][13]],
                        ['Dogrib(Tlicho)',cma_data['DATA'][416][13]],
                        ['Sekani',cma_data['DATA'][419][13]],
                        ['NorthSlavey(Hare)',cma_data['DATA'][421][13]],
                        ['SouthSlavey',cma_data['DATA'][422][13]],
                        ['Slaveyn.o.s.',cma_data['DATA'][423][13]],
                        ['Kaska(Nahani)',cma_data['DATA'][425][13]],
                        ['Tahltan',cma_data['DATA'][426][13]],
                        ['NorthernTutchone',cma_data['DATA'][428][13]],
                        ['SouthernTutchone',cma_data['DATA'][429][13]],
                        ['Haida',cma_data['DATA'][431][13]],
                        ['Inuinnaqtun(Inuvialuktun)',cma_data['DATA'][433][13]],
                        ['Kutenai',cma_data['DATA'][441][13]],
                        ['Comox',cma_data['DATA'][444][13]],
                        ['Halkomelem',cma_data['DATA'][445][13]],
                        ['Lillooet',cma_data['DATA'][446][13]],
                        ['Okanagan',cma_data['DATA'][447][13]],
                        ['Shuswap(Secwepemctsin)',cma_data['DATA'][448][13]],
                        ['Squamish',cma_data['DATA'][449][13]],
                        ['Straits',cma_data['DATA'][450][13]],
                        ['Thompson(Ntlakapamux)',cma_data['DATA'][451][13]],
                        ['Stoney',cma_data['DATA'][455][13]],
                        ['Tlingit',cma_data['DATA'][457][13]],
                        ['Gitxsan(Gitksan)',cma_data['DATA'][459][13]],
                        ['NisgaaCHANGESPELLING',cma_data['DATA'][460][13]],
                        ['Tsimshian',cma_data['DATA'][461][13]],
                        ['Haisla',cma_data['DATA'][463][13]],
                        ['Heiltsuk',cma_data['DATA'][464][13]],
                        ['Kwakiutl(KwakwalaCHANGESPELLING)',cma_data['DATA'][465][13]]]
        language_groups.sort(key= lambda x:x[1], reverse=True)
        seniors_count = cma_data['DATA'][36][13]
        median_age = cma_data['DATA'][39][13]

        # Get Prov 2016 Census Data
        prov_id = '2016A000235'
        with urllib.request.urlopen('https://www12.statcan.gc.ca/rest/census-recensement/CPR2016.json?lang=E&dguid='+prov_id) as url:
                prov_data = json.loads(url.read().decode())
        total_pop_count_prov = prov_data['DATA'][0][13]
        immigrants_count_prov = prov_data['DATA'][1141][13]
        immigrants_2011to2016_count_prov = prov_data['DATA'][1148][13]
        non_citizen_count_prov = prov_data['DATA'][1138][13]
        industry_total_prov =  prov_data['DATA'][1898][13]
        industry_groups_prov = {
                        'agriculture, forestry, fishing and hunting':prov_data['DATA'][1899][13],
                        'mining, quarrying, and oil and gas extraction':prov_data['DATA'][1900][13],
                        'utilities':prov_data['DATA'][1901][13],
                        'construction':prov_data['DATA'][1902][13],
                        'manufacturing':prov_data['DATA'][1903][13],
                        'wholesale trade':prov_data['DATA'][1904][13],
                        'retail trade':prov_data['DATA'][1905][13],
                        'transportation and warehousing':prov_data['DATA'][1906][13],
                        'information and cultural industries':prov_data['DATA'][1907][13],
                        'finance and insurance':prov_data['DATA'][1908][13],
                        'real estate and rental and leasing':prov_data['DATA'][1909][13],
                        'professional, scientific and technical services':prov_data['DATA'][1910][13],
                        'management of companies and enterprises':prov_data['DATA'][1911][13],
                        'administrative and support, waste management and remediation services':prov_data['DATA'][1912][13],
                        'educational services':prov_data['DATA'][1913][13],
                        'health care and social assistance':prov_data['DATA'][1914][13],
                        'arts, entertainment and recreation':prov_data['DATA'][1915][13],
                        'accommodation and food services':prov_data['DATA'][1916][13],
                        'other services (except public administration)':prov_data['DATA'][1917][13],
                        'public administration':prov_data['DATA'][1918][13]
                        }

        #####################################################################
        ######################LFS Data for the CMA###########################
        #####################################################################
                
            
        lfs_cma = StatcanTable(14100380)
    
        #Filtering data for ON CMAs only, estimates and seasonally adjusted 
        lfs_cma = lfs_cma[lfs_cma['GEO'].str.contains("Ontario")]
        lfs_cma = lfs_cma.loc[(lfs_cma['Statistics']=='Estimate') & (lfs_cma['Data type']=='Seasonally adjusted')]
            
        #Filtering lfs_cma to get data only for selected CMA:

        lfs_cma = lfs_cma.loc[(lfs_cma['DGUID']==cma_id)]
        #changing the date format for REF_DATE column 
        lfs_cma['REF_DATE'] = pd.to_datetime(lfs_cma['REF_DATE'])
        ly = max(lfs_cma['REF_DATE']) - relativedelta(years=1)
        month = max(lfs_cma['REF_DATE'])
        #Filtering CMA data to get ON unemployment rate, participation rate and employment rate
        cma_labour_force = lfs_cma.loc[(lfs_cma['Labour force characteristics']=="Labour force") & (lfs_cma['REF_DATE']==max(lfs_cma['REF_DATE'])),'VALUE'].iloc[0]
        cma_unemp_rate = lfs_cma.loc[(lfs_cma['Labour force characteristics']=="Unemployment rate") & (lfs_cma['REF_DATE']==max(lfs_cma['REF_DATE'])),'VALUE'].iloc[0]
        cma_part_rate = lfs_cma.loc[(lfs_cma['Labour force characteristics']=="Participation rate") & (lfs_cma['REF_DATE']==max(lfs_cma['REF_DATE'])),'VALUE'].iloc[0]
        cma_emp_rate = lfs_cma.loc[(lfs_cma['Labour force characteristics']=="Employment rate") & (lfs_cma['REF_DATE']==max(lfs_cma['REF_DATE'])),'VALUE'].iloc[0]
        cma_unemp_rate_change = lfs_cma.loc[(lfs_cma['Labour force characteristics']=="Unemployment rate") & (lfs_cma['REF_DATE']==ly),'VALUE'].iloc[0] \
        - cma_unemp_rate
        cma_emp_change = lfs_cma.loc[(lfs_cma['Labour force characteristics']=="Employment") & (lfs_cma['REF_DATE']==max(lfs_cma['REF_DATE'])),'VALUE'].iloc[0] \
        - lfs_cma.loc[(lfs_cma['Labour force characteristics']=="Employment") & (lfs_cma['REF_DATE']==ly),'VALUE'].iloc[0]



        #### INCOME FOR CMAs, zach ####

        # 1) Go to https://www12.statcan.gc.ca/census-recensement/2021/dp-pd/prof/details/download-telecharger.cfm?Lang=E #
        # 2) Download Census metropolitan areas (CMAs) and census agglomerations (CAs), CSV #
        # 3) Place in same folder as .py file #
        # 4) Change first line below to directory you are using #
    
        Census21 = pd.read_csv (r"/Users/zack/Documents/Python/Community Backgrounders/98-401-X2021002_eng_CSV/98-401-X2021002_English_CSV_data.csv",low_memory=False, encoding='latin-1')
        cma_id_2 = cma_id_list_2021[i] 
        Census_cma = Census21.loc[Census21['DGUID']==cma_id_2]
        cma_med_tot_inc = Census_cma.loc[Census_cma['CHARACTERISTIC_NAME'].str.contains("Median total income of household in 2020"),'C1_COUNT_TOTAL'].iloc[0]
        cma_prevalence_lim = Census_cma.loc[Census_cma['CHARACTERISTIC_NAME'].str.contains("Prevalence of low income based on the Low-income measure, after tax"),'C1_COUNT_TOTAL'].iloc[0]        
        
        # Define variables for Ontario's rates 
        lfs_on = StatcanTable(14100287)
        lfs_on = lfs_on.loc[(lfs_on['Statistics']=='Estimate') & (lfs_on['Data type']=='Seasonally adjusted')\
                           & (lfs_on['GEO']=='Ontario') & (lfs_on['Sex']=='Both sexes') & (lfs_on['Age group']=='15 years and over')]
        #changing the date format for REF_DATE column
        lfs_on['REF_DATE'] = pd.to_datetime(lfs_on['REF_DATE'])
        ly = max(lfs_on['REF_DATE']) - relativedelta(years=1)
        
        on_unemp_rate = lfs_on.loc[(lfs_on['Labour force characteristics']=="Unemployment rate") & (lfs_on['REF_DATE']==max(lfs_on['REF_DATE'])),'VALUE'].iloc[0]
        on_part_rate = lfs_on.loc[(lfs_on['Labour force characteristics']=="Participation rate") & (lfs_on['REF_DATE']==max(lfs_on['REF_DATE'])),'VALUE'].iloc[0]
        on_emp_rate = lfs_on.loc[(lfs_on['Labour force characteristics']=="Employment rate") & (lfs_on['REF_DATE']==max(lfs_on['REF_DATE'])),'VALUE'].iloc[0]
        on_unemp_rate_change = lfs_on.loc[(lfs_on['Labour force characteristics']=="Unemployment rate") & (lfs_on['REF_DATE']==ly),'VALUE'].iloc[0] \
        - on_unemp_rate
        
        
        #################################################################
        ############## People with Disabilities Text ####################
        #################################################################
        
        if cma_id == '2016S050535505': 
            cma_id ='2016A000235'
        pwd = StatcanTable(13100750)
        pwd = pwd.loc[pwd['DGUID']==cma_id]
        cma_pwd_num = pwd.loc[(pwd['Disability']=='Persons with disabilities') & (pwd['Estimates'] == 'Number of persons'), 'VALUE'].iloc[0]
        cma_pwd_pct = pwd.loc[(pwd['Disability']=='Persons with disabilities') & (pwd['Estimates'] == 'Percentage of persons'), 'VALUE'].iloc[0]

        
        ####################################################
        # 2021 CMA Census Data
        
        
        cma_id = cma_id_list_2021[i]
        
        #Getting 2021 Census Population count data for Ontario and CMAs - Table 98-10-0005
        df = StatcanTable(98100005)
        df.rename(columns = {'Population and dwelling counts (13): Population, 2021 [1]':'POP_2021','Population and dwelling counts (13): Population, 2016 [2]':'POP_2016','Population and dwelling counts (13): Population percentage change, 2016 to 2021 [3]':'POP_CHANGE_16_21'},inplace=True)
        # Keeping only the columns we want
        df = df.iloc[:,[0,1,2,4,6,8]]
        on_total_pop_count = df.loc[df['DGUID']=='2021A000235','POP_2021'].iloc[0]
        on_pop_pct_change_16_21 = df.loc[df['DGUID']=='2021A000235','POP_CHANGE_16_21'].iloc[0]
        total_pop_count = df.loc[df['DGUID']==cma_id,'POP_2021'].iloc[0]
        pop_pct_change_16_21 = (df.loc[df['DGUID']==cma_id,'POP_CHANGE_16_21'].iloc[0])
            
        #Getting 2021 Census Seniors Population numbers for Ontario
        df = StatcanTable(98100025)
        df.columns = ['REF_DATE','GEO','DGUID','AGE','COORDINATE','GENDER-TOTAL','SYMBOL','GENDER-MEN+','SYMBOL1','GENDER-WOMEN+','SYMBOL2']
            
        on_med_age_21 = df.loc[(df['DGUID']=='2021A000235') & (df['AGE']=="Median age"),'GENDER-TOTAL'].iloc[0]
        on_seniors_pop_21 = df.loc[(df['DGUID']=='2021A000235') & (df['AGE']=="65 years and over"),'GENDER-TOTAL'].iloc[0]
        on_total_pop_21 = df.loc[(df['DGUID']=='2021A000235') & (df['AGE']=="Total - Age"),'GENDER-TOTAL'].iloc[0]
        on_seniors_prop_21 = round(on_seniors_pop_21/on_total_pop_21 * 100 ,1)
            
        #Getting 2021 Census Seniors Population numbers for CMA

        df = StatcanTable(98100024)
        df.columns = ['REF_DATE','GEO','DGUID','AGE','COORDINATE','GENDER-TOTAL','SYMBOL','GENDER-MEN+','SYMBOL1','GENDER-WOMEN+','SYMBOL2']
                 
        cma_med_age_21 = df.loc[(df['DGUID']==cma_id) & (df['AGE']=="Median age"),'GENDER-TOTAL'].iloc[0]
        seniors_pop_21 = df.loc[(df['DGUID']==cma_id) & (df['AGE']=="65 years and over"),'GENDER-TOTAL'].iloc[0]
        total_pop_21 = df.loc[(df['DGUID']==cma_id) & (df['AGE']=="Total - Age"),'GENDER-TOTAL'].iloc[0]
        seniors_prop_21 = round(seniors_pop_21/total_pop_21 * 100,1)

        
        
        ####################################################
        # Generate DOCX Document
        document = Document()
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
        
        ############################################################################↓↓↓↓↓FORMATTING
        style = document.styles["Normal"]
        font = style.font
        font.name = 'Arial'

        style = document.styles['Heading 4']
        font = style.font
        rFonts = font.element.rPr.rFonts
        rFonts.set(qn("w:asciiTheme"), "Arial")
        font.name = 'Arial'
        font.size = Pt(28)
        font.color.rgb = RGBColor(192,0,0)
        font.italic = False
        font.bold = False

        style = document.styles['Title']
        font = style.font
        rFonts = font.element.rPr.rFonts
        rFonts.set(qn("w:asciiTheme"), "Arial")
        font.name = 'Arial'
        font.size = Pt(14)
        font.color.rgb = RGBColor(192,0,0)
        font.italic = False
        font.bold = True

        style = document.styles['Heading 3']
        font = style.font
        rFonts = font.element.rPr.rFonts
        rFonts.set(qn("w:asciiTheme"), "Arial")
        font.name = 'Arial'
        font.size = Pt(11)
        font.color.rgb = RGBColor(192,0,0)
        font.italic = False
        font.bold = True

        par1=document.add_paragraph(f'Profile:',style="Heading 4")
        par1.style = document.styles['Heading 4']

        para=document.add_paragraph("Constituency:",style="Title")
        ############################################################################↑↑↑↑↑FORMATTING
        
        document.add_paragraph('Political Affiliation:',style="Heading 3")
        document.add_paragraph('Parliamentary Tenure:',style="Heading 3")
        document.add_paragraph('Notable Parliametary Roles:',style="Heading 3")
        document.add_paragraph('Background:',style="Heading 3")
        document.add_paragraph('Media Scan:',style="Heading 3")
        p=document.add_paragraph('Constituency Demographics:',style="Heading 3").add_run()
        p.add_break()
        p=document.add_paragraph().add_run('Geography')
        p.bold=True
        p=document.add_paragraph().add_run('Population')
        p.bold=True
        
        document.add_paragraph(
                        f'According to the 2021 Census, the population of the {cma_name} CMA was {total_pop_count:,.0f}, '+
                        ('an increase ' if pop_pct_change_16_21>0 else 'a decrease ')+
                        f'of {pop_pct_change_16_21/100:.1%} from 2016.  '+
                        ('This is in contrast to the province where the population grew by 775,448 (5.8%).' if pop_pct_change_16_21 != 5.8 else
                        'The population of the province grew by 775,448, which was also an increase of 5.8%.'))
        
        
        p=document.add_paragraph().add_run('Indigenous')
        p.bold=True
        document.add_paragraph(
                        f'In 2016, {indigenous_count:,.0f} individuals in the {cma_name} CMA self-identified as Indigenous, '+
                        f'representing about {indigenous_count/total_pop_count:.1%} '+
                        f'of the population. In comparison, 2.8% (374,395) of the provincial population self-identified as Indigenous.')

        p=document.add_paragraph().add_run('Visible Minorities')
        p.bold=True
        document.add_paragraph(
                        f'Roughly {vis_minority_count/vis_minority_total:.1%} of the {cma_name} CMA population self-identified as a '+
                        f'visible minority in 2016, or {vis_minority_count:,.0f}  '+
                        f'individuals. '+
                        f'{vis_minority_groups[0][0]} ({vis_minority_groups[0][1]:,.0f}; {vis_minority_groups[0][1]/vis_minority_total:.1%}), '+
                        f'{vis_minority_groups[1][0]} ({vis_minority_groups[1][1]:,.0f}; {vis_minority_groups[1][1]/vis_minority_total:.1%}), '+
                        f'and '+
                        f'{vis_minority_groups[2][0]} ({vis_minority_groups[2][1]:,.0f}; {vis_minority_groups[2][1]/vis_minority_total:.1%}) '+
                        f'were the top three ethnicities. About 29.3% or 3,885,585 individuals in Ontario identified  '+
                        f'as a visible minority, and the top three ethnicities were South Asian (1,150,415; 8.7%), Chinese (754,550; 5.7%) and Black (627,715; 4.7%).')

        p=document.add_paragraph().add_run('Persons with Disabilities')
        p.bold=True
        document.add_paragraph(
                        f'While data for the {cma_name} cma is scarce, {cma_pwd_num:,.0f} ({cma_pwd_pct}%) of the {cma_name} CMA population identified as persons '+
                        f'with disabilities in 2017. Comparatively, 2,616,170 '+
                        f'(24.1%) individuals identified as persons with disabilities at the provincial level.')
        
        
        p=document.add_paragraph().add_run('Newcomers')
        p.bold=True
        document.add_paragraph(
                        f'About {immigrants_count/immigrants_total:.1%} ({immigrants_count:,.0f}) of the {cma_name} CMA’s population were '+
                        f'immigrants at the time of the 2016 Census. About '+
                        f'{immigrants_2011to2016_count/immigrants_total:.1%} ({immigrants_2011to2016_count:,.0f}) immigrated between '+
                        f'2011 and 2016. In comparison, 29.1% (3,852,145) of Ontario’s '+
                        f'population were immigrants at the time of the Census and 3.6% (472,170) immigrated into the province between '+
                        f'2011 and 2016. Roughly '+
                        f'{non_citizen_count/immigrants_total:.1%} ({non_citizen_count:,.0f}) of the {cma_name} CMA’s population and '+
                        f'7.7% (1,019,095) of Ontario’s population did not have a '+
                        f'Canadian Citizenship at the time of the 2016 Census.')

        p=document.add_paragraph().add_run('Languages')
        p.bold=True
        document.add_paragraph(
                        f'Approximately {englishandfrench/language_total:.1%} ({englishandfrench:,.0f}) of the {cma_name} CMA '+
                        f'population had knowledge of both official languages at the time '+
                        f'of the 2016 Census, compared to 11.2% (1,490,390) of the province. Optional sentence: ____ is an official '+
                        f'language minority community with at least '+
                        f'one school where the primary teaching language is the official language minority language. '+
                        f'Apart from official languages, the top three languages most '+
                        f'spoken at home in the CMA were '+
                        f'{language_groups[0][0]} ({language_groups[0][1]:,.0f}; {language_groups[0][1]/language_total:.1%}), '+
                        f'{language_groups[1][0]} ({language_groups[1][1]:,.0f}; {language_groups[1][1]/language_total:.1%}), and '+
                        f'{language_groups[2][0]} ({language_groups[2][1]:,.0f}; {language_groups[2][1]/language_total:.1%}) '+
                        f'Comparatively, the top three languages most spoken at home in Ontario were '+
                        f'Mandarin (220,535; 1.7%), Cantonese (198,745; 1.5%), and Punjabi (132,135; 1%).')

        #FROM LAST SAVE:
        #####################################################################################################
        
        # VARIABLES FOR SENIORS SECTION
        
        seniors_prop_21 = round(seniors_pop_21/total_pop_21 * 100,1)
        on_seniors_prop_21 = round(on_seniors_pop_21/on_total_pop_21 * 100 ,1)

        if cma_med_age_21 > on_med_age_21: 
            choice_1='an older '
        elif cma_med_age_21 < on_med_age_21:
            choice_1='a younger '
        elif cma_med_age_21 == on_med_age_21:
            choice_1='a similar '

        if seniors_prop_21<(on_seniors_prop_21-0.2):
            choice_2=f'There were proportionally fewer seniors, with {seniors_prop_21}% of the population 65 years or older, versus {on_seniors_prop_21}% for the province.'
        elif (on_seniors_prop_21-0.2)<=seniors_prop_21<=(on_seniors_prop_21+0.2):
            choice_2=f"The proportion of seniors was similar to that of the province. About {seniors_prop_21}% of the {cma_name} CMA's population was 65 years or older, versus {on_seniors_prop_21}% for all of Ontario."
        elif seniors_prop_21>(on_seniors_prop_21+0.2):
            choice_2=f'There were proportionally more seniors, with {seniors_prop_21}% of the population 65 years or older, versus {on_seniors_prop_21}% for the province.'
        
        #CREATING SENIORS SECTION PARAGRAPH
        p=document.add_paragraph().add_run('Seniors')
        p.bold=True
        
        document.add_paragraph(
                        f'The {cma_name} CMA had '+
                        str(choice_1)+
                        f'age profile compared to the province. According to the 2021 Census, the median age was '+
                        str(cma_med_age_21)+
                        f', compared with {on_med_age_21} for Ontario. '+
                        choice_2)
        #END SENIORS SECTION
        
        
        p=document.add_paragraph().add_run('Labour Force')
        p.bold=True
        document.add_paragraph(
                        f'According to the Labour Force Survey, the {cma_name} CMA had a labour force of about {(cma_labour_force *1000):,.0f} in {month} 2022. '+
                        f'The unemployment rate was {cma_unemp_rate}% —'+ ('higher' if (cma_unemp_rate>on_unemp_rate) else 'lower') +
                        f' than Ontario’s {on_unemp_rate}%.' +
                        f'The participation rate and the employment rate in the CMA were ' + ('higher' if (cma_emp_rate>on_emp_rate) else 'lower') +
                        f' than the province, '+
                        f'at {cma_emp_rate}% and {cma_part_rate}% respectively. By comparison, Ontario’s participation rate was {on_part_rate}%, '+
                        f'and the employment rate was {on_emp_rate}%. '+
                        f'Labour market conditions in the {cma_name} CMA ' + ('improved' if (cma_emp_change>0) else 'declined') +
                        f' between {month} 2021 and {month} 2022 as employment '+
                        ('grew' if (cma_emp_change >0) else 'shrank') +
                        f' by {abs(cma_emp_change*1000):,.0f}. The unemployment rate in the {cma_name} CMA ' +
                        ('increased' if (cma_unemp_rate_change>0) else 'decreased') +
                        f' by {abs(cma_unemp_rate_change):,.1f} percentage points to {cma_unemp_rate}% over the year, '+
                        f'compared to a' + ('n increase' if (on_unemp_rate_change>0) else 'decrease') +
                        f'of {abs(on_unemp_rate_change)} percentage points provincially.')

        #### Income Paragraph for CMA, zach ####
        p=document.add_paragraph().add_run('Income')
        p.bold=True
        document.add_paragraph(
                        f'The median total income for the {cma_name} CMA private households in 2020 was ${cma_med_tot_inc:,.0f} which was '+
                        ('higher ' if  (cma_med_tot_inc>91000) else 'lower ')+
                        f'than Ontario ($91,000). '+
                        f'The prevalence of low-income based on the Low-income measure after-tax (LIM-AT) indicator is {cma_prevalence_lim}% '+
                        f'in the {cma_name} CMA compared to 10.1% for Ontario.')  


        p=document.add_paragraph().add_run('Industrial Base')
        p.bold=True
        over_rep_industries = []
        for x in industry_groups:
                x_industry_name = x[0]
                x_prov_pct = industry_groups_prov[x_industry_name]/industry_total_prov
                x_cma_pct = x[1]/industry_total
                if x_cma_pct/x_prov_pct>=2:
                        over_rep_industries.append([x_industry_name,x_prov_pct,x_cma_pct])
        over_rep_industries.sort(key= lambda x:x[2]/x[1], reverse=True)
        document.add_paragraph(
                        f'According to the 2016 Census, the largest three sectors by employment size in the {cma_name} CMA were '+
                        f'{industry_groups[0][0]} ({industry_groups[0][1]:,.0f}; {industry_groups[0][1]/industry_total:.1%}), '+
                        f'{industry_groups[1][0]} ({industry_groups[1][1]:,.0f}; {industry_groups[1][1]/industry_total:.1%}), and '+
                        f'{industry_groups[2][0]} ({industry_groups[2][1]:,.0f}; {industry_groups[2][1]/industry_total:.1%}). '+
                        f'Employment in these top three sectors made up about '+
                        f'{industry_groups[0][1]/industry_total+industry_groups[1][1]/industry_total+industry_groups[2][1]/industry_total:.1%} '+
                        f'of the employed labour force in the {cma_name} CMA, '+
                        ('higher' if ((industry_groups[0][1]+industry_groups[1][1]+industry_groups[2][1])/industry_total>(industry_groups_prov[industry_groups[0][0]]+industry_groups_prov[industry_groups[1][0]]+industry_groups_prov[industry_groups[2][0]])/industry_total_prov)
                         else 'lower')+
                        f' than the Ontario share of '+
                        f'{industry_groups_prov[industry_groups[0][0]]/industry_total_prov+industry_groups_prov[industry_groups[1][0]]/industry_total_prov+industry_groups_prov[industry_groups[2][0]]/industry_total_prov:.1%}. '+
                        ('Additionally, ' if len(over_rep_industries)>0 else '')+ 
                        (', '.join([x[0] for x in over_rep_industries[:-1]])+', and '+over_rep_industries[-1][0] if len(over_rep_industries)>2 else '')+
                        (over_rep_industries[0][0]+' and '+over_rep_industries[1][0] if len(over_rep_industries)==2 else '')+
                        (over_rep_industries[0][0] if len(over_rep_industries)==1 else '')+
                        (' is ' if len(over_rep_industries)==1 else ' are ' if len(over_rep_industries)>1 else '')+
                        (f'overrepresented in the {cma_name} CMA. ' if len(over_rep_industries)>0 else '')+
                        (''.join([x[0].capitalize()+f' represents about {x[2]:.1%} of the employment in the CMA compared to {x[1]:.1%} at the provincial level. ' for x in over_rep_industries])))

        
        document.save(cma_name+'_CMA.docx')
        print (f'{cma_name} is done')




# In[ ]:


pwd = StatcanTable(13100750)
pwd = pwd.loc[pwd['GEO'].str.contains('Ottawa')]
pwd
#cma_pwd_num = pwd.loc[(pwd['Disability']=='Persons with disabilities') & (pwd['Estimates'] == 'Number of persons'), 'VALUE'].iloc[0]
#cma_pwd_pct = pwd.loc[(pwd['Disability']=='Persons with disabilities') & (pwd['Estimates'] == 'Percentage of persons'), 'VALUE'].iloc[0]

